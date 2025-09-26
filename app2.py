"""
Agentic Resume Tailor (Streamlit)

Streamlit app that tailors resumes using LangChain + Google Gemini.
User uploads resume + JD, app suggests edits, and exports final DOCX/PDF.
"""

import streamlit as st
import os
import io
import json
import re
from typing import List, Tuple
from dotenv import load_dotenv
import numpy as np

import PyPDF2
import docx
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# Load env key once (backend only, never shown to user)
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

try:
    from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
except Exception:
    ChatGoogleGenerativeAI = None
    GoogleGenerativeAIEmbeddings = None


# -------------------- Helpers: document parsing --------------------

def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    except Exception:
        return ""


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    try:
        bio = io.BytesIO(docx_bytes)
        doc = docx.Document(bio)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception:
        return ""


def extract_text_from_uploaded(uploaded_file) -> Tuple[str, str]:
    if not uploaded_file:
        return "", ""
    name = uploaded_file.name
    raw = uploaded_file.read()
    lname = name.lower()
    if lname.endswith(".pdf"):
        return extract_text_from_pdf_bytes(raw), name
    if lname.endswith(".docx"):
        return extract_text_from_docx_bytes(raw), name
    if lname.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore"), name
    return raw.decode("utf-8", errors="ignore"), name


# -------------------- LLM + embeddings --------------------

def init_llm_and_embeddings():
    if not GEMINI_API_KEY:
        raise RuntimeError("No GEMINI_API_KEY found. Add it to .env or environment variables.")

    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        temperature=0,
        google_api_key=GEMINI_API_KEY
    )
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001",
        google_api_key=GEMINI_API_KEY
    )
    return llm, embeddings


def safe_json_extract(s: str):
    s = s.strip()
    try:
        return json.loads(s)
    except Exception:
        pass
    m = re.search(r"(\[.*?\]|\{.*?\})", s, flags=re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception:
            pass
    return [l.strip(" -•") for l in s.splitlines() if l.strip()]


def extract_keywords_from_jd(llm, job_text: str, max_keywords: int = 50) -> List[str]:
    prompt = (
        f"Extract up to {max_keywords} important keywords from this job description. "
        "Reply with a JSON array of lowercase keywords only.\n\n" + job_text
    )
    ai_msg = llm.invoke([("system", "Keyword extractor."), ("human", prompt)])
    parsed = safe_json_extract(getattr(ai_msg, "content", str(ai_msg)))
    return [str(x).lower().strip() for x in parsed if str(x).strip()][:max_keywords]


def extract_resume_bullets(llm, resume_text: str, max_bullets: int = 40) -> List[str]:
    prompt = (
        f"Extract up to {max_bullets} resume bullet points (achievements, responsibilities). "
        "Return JSON array only.\n\n" + resume_text
    )
    ai_msg = llm.invoke([("system", "Resume bullet extractor."), ("human", prompt)])
    parsed = safe_json_extract(getattr(ai_msg, "content", str(ai_msg)))
    return [str(x).strip() for x in parsed if str(x).strip()][:max_bullets]


def reword_variations_for_bullet(llm, bullet: str, n: int = 3) -> List[str]:
    prompt = f"Rewrite this resume bullet in {n} different, concise ways. JSON array only.\n\n{bullet}"
    ai_msg = llm.invoke([("system", "Resume rewriter."), ("human", prompt)])
    parsed = safe_json_extract(getattr(ai_msg, "content", str(ai_msg)))
    return [str(x).strip() for x in parsed if str(x).strip()][:n]


# -------------------- UI --------------------

st.set_page_config(page_title="Agentic Resume Tailor", layout="wide")
st.title("Agentic Resume Tailor — Streamlit (LangChain + Gemini)")

with st.sidebar:
    show_raw = st.checkbox("Show raw extracted text", value=False)
    sample_mode = st.checkbox("Use fast demo mode", value=False)

col1, col2 = st.columns(2)
with col1:
    resume_file = st.file_uploader("Upload resume", type=["pdf", "docx", "txt"])
with col2:
    jd_file = st.file_uploader("Upload job description", type=["pdf", "docx", "txt"])

if st.button("Analyze"):
    if not resume_file or not jd_file:
        st.error("Please upload both a resume and a job description.")
        st.stop()

    llm, embeddings = init_llm_and_embeddings()
    resume_text, _ = extract_text_from_uploaded(resume_file)
    jd_text, _ = extract_text_from_uploaded(jd_file)

    if show_raw:
        st.subheader("Raw texts")
        st.text_area("Resume", resume_text, height=200)
        st.text_area("Job Description", jd_text, height=200)

    jd_keywords = extract_keywords_from_jd(llm, jd_text, 50)
    resume_bullets = extract_resume_bullets(llm, resume_text, 60)

    st.success(f"Extracted {len(jd_keywords)} keywords and {len(resume_bullets)} resume bullets")

    # Tailored bullet suggestions
    st.subheader("Tailored bullet suggestions")
    bullets_to_show = resume_bullets if not sample_mode else resume_bullets[:6]

    st.session_state["new_bullets"] = []
    for i, b in enumerate(bullets_to_show):
        st.markdown(f"**Original bullet {i+1}:** {b}")
        try:
            variations = reword_variations_for_bullet(llm, b, 3)
        except Exception:
            variations = []
        choice = st.selectbox(
            f"Variation for bullet {i+1}",
            ["(keep original)"] + variations,
            key=f"var_{i}"
        )
        edited = st.text_area(f"Edit bullet {i+1}", value=b if choice == "(keep original)" else choice, key=f"edit_{i}")
        st.session_state["new_bullets"].append(edited)

    # Finalize + export
    st.subheader("Finalize & export")
    full_name = st.text_input("Name (header)")
    contact = st.text_input("Contact info")

    if st.button("Generate final DOCX and PDF"):
        final_bullets = st.session_state.get("new_bullets", [])
        if not final_bullets:
            st.error("No bullets available.")
        else:
            # DOCX
            docx_io = io.BytesIO()
            doc = docx.Document()
            if full_name: doc.add_heading(full_name, level=1)
            if contact: doc.add_paragraph(contact)
            doc.add_heading("Experience", level=2)
            for b in final_bullets:
                doc.add_paragraph(b, style="List Bullet")
            doc.save(docx_io)
            docx_io.seek(0)

            # PDF
            pdf_io = io.BytesIO()
            elements = []
            styles = getSampleStyleSheet()
            if full_name: elements.append(Paragraph(full_name, styles["Title"]))
            if contact: elements.append(Paragraph(contact, styles["Normal"]))
            elements.append(Spacer(1, 12))
            elements.append(Paragraph("Experience", styles["Heading2"]))
            for b in final_bullets:
                elements.append(Paragraph(b, styles["Normal"]))
                elements.append(Spacer(1, 6))
            SimpleDocTemplate(pdf_io, pagesize=letter).build(elements)
            pdf_io.seek(0)

            # Downloads
            st.download_button("Download DOCX", docx_io.getvalue(), "tailored_resume.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.download_button("Download PDF", pdf_io.getvalue(), "tailored_resume.pdf", mime="application/pdf")
            st.success("Resume exported successfully!")
