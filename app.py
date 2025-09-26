"""
Agentic Resume Tailor (Streamlit)

Single-file Streamlit app that demonstrates a simple "agentic" resume tailor pipeline
using LangChain (langchain-google-genai) + Google Gemini (Gemini API) for LLM calls and
embeddings, and PyPDF2 / python-docx for document handling.

Features implemented (minimal working prototype):
- Upload resume and job-description (PDF/DOCX/TXT)
- Extract keywords from job description (LLM)
- Extract bullets/experience from resume (LLM)
- Exact + semantic matching (embeddings) to find matched/missing keywords
- Color-coded alignment score
- Generate 2-3 rewording variations per bullet (LLM)
- Let user pick/edit suggestions in-app
- Export final resume as DOCX and PDF

NOTE / RUN:
1) Install dependencies (suggested):
    pip install -r requirements.txt

requirements.txt (recommended):
# core
streamlit
numpy
PyPDF2
python-docx
reportlab
# LangChain + Gemini integration
langchain
langchain-google-genai
google-genai

2) Set your Google AI API key in an env var (or paste inside the app sidebar):
    export GOOGLE_API_KEY="ya..."

3) Run app:
    streamlit run agentic_resume_tailor_app.py

IMPORTANT: This is a starting prototype. API/model names and integration packages may change
over time (see LangChain's docs for "langchain-google-genai" and Google's Gemini API).
"""

import streamlit as st
import os
import io
import json
import re
import tempfile
import base64
from typing import List, Tuple, Dict, Any
from dotenv import load_dotenv
import numpy as np

import PyPDF2
import docx
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# LangChain / Google Gemini integration
# Uses the langchain-google-genai wrapper (langchain_google_genai)
# and the ChatGoogleGenerativeAI + GoogleGenerativeAIEmbeddings classes.
load_dotenv()
gemini_api_key = os.getenv("GEMINI_API_KEY")
try:
    from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
except Exception:
    # graceful fallback so the app can show a helpful message in the UI
    ChatGoogleGenerativeAI = None
    GoogleGenerativeAIEmbeddings = None

# -------------------- Helpers: document parsing --------------------

def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        texts = []
        for p in reader.pages:
            t = p.extract_text()
            if t:
                texts.append(t)
        return "\n".join(texts)
    except Exception as e:
        return ""  # fallback empty


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    try:
        bio = io.BytesIO(docx_bytes)
        doc = docx.Document(bio)
        paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(paras)
    except Exception as e:
        return ""


def extract_text_from_uploaded(uploaded_file) -> Tuple[str, str]:
    """Return (text, filename)
    Supports PDF, DOCX, TXT. If unsupported, returns raw bytes decoded if possible.
    """
    if not uploaded_file:
        return "", ""

    name = uploaded_file.name
    raw = uploaded_file.read()
    lname = name.lower()
    if lname.endswith('.pdf'):
        return extract_text_from_pdf_bytes(raw), name
    if lname.endswith('.docx'):
        return extract_text_from_docx_bytes(raw), name
    if lname.endswith('.txt') or 'text' in uploaded_file.type:
        try:
            return raw.decode('utf-8', errors='ignore'), name
        except Exception:
            return str(raw), name
    # fallback: try pdf and docx extraction
    txt = extract_text_from_pdf_bytes(raw)
    if txt.strip():
        return txt, name
    txt = extract_text_from_docx_bytes(raw)
    if txt.strip():
        return txt, name
    try:
        return raw.decode('utf-8', errors='ignore'), name
    except Exception:
        return str(raw), name

# -------------------- Helpers: LLM + embeddings --------------------

def init_llm_and_embeddings(api_key: str):
    if not api_key:
        api_key = gemini_api_key

    if not api_key:
        raise RuntimeError('No GOOGLE_API_KEY provided. Set env var or paste it in the sidebar.')

    if ChatGoogleGenerativeAI is None or GoogleGenerativeAIEmbeddings is None:
        raise RuntimeError('langchain-google-genai package not installed. See requirements.')

    # ✅ pass the API key explicitly
    llm = ChatGoogleGenerativeAI(
        model='gemini-2.5-flash',
        temperature=0,
        google_api_key=api_key
    )

    embeddings = GoogleGenerativeAIEmbeddings(
        model='models/embedding-001',
        google_api_key=api_key
    )

    return llm, embeddings



def safe_json_extract(s: str):
    """Try to find and parse JSON inside a model output. Returns parsed object or None."""
    s = s.strip()
    # direct parse
    try:
        return json.loads(s)
    except Exception:
        pass
    # try to extract the first JSON array or object in the text
    m = re.search(r'({.*?})', s, flags=re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception:
            pass
    m = re.search(r'(\[.*?\])', s, flags=re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception:
            pass
    # last resort: try to line-split into tokens
    lines = [l.strip().strip('-• ') for l in s.splitlines() if l.strip()]
    return lines

# -------------------- Domain prompts & LLM helpers --------------------

def extract_keywords_from_jd(llm, job_text: str, max_keywords: int = 50) -> List[str]:
    prompt = (
        "Extract up to %d important keywords (skills, technologies, certifications, role-specific terms) "
        "from the following job description. Reply with a JSON array of lowercase keywords (no explanations).\n\n"
        "Job description:\n" % (max_keywords)
    ) + job_text

    messages = [
        ("system", "You are an assistant that extracts concise, unambiguous keywords from job descriptions."),
        ("human", prompt),
    ]
    ai_msg = llm.invoke(messages)
    content = getattr(ai_msg, 'content', str(ai_msg))
    parsed = safe_json_extract(content)
    # normalize into a list of strings
    if isinstance(parsed, list):
        return [str(x).lower().strip() for x in parsed if str(x).strip()]
    if isinstance(parsed, dict):
        # if returned as {"keywords": [...]}
        for v in parsed.values():
            if isinstance(v, list):
                return [str(x).lower().strip() for x in v if str(x).strip()]
    # fallback: split lines
    if isinstance(parsed, str):
        return [k.strip().lower() for k in parsed.split(',') if k.strip()]
    if isinstance(parsed, (tuple, set)):
        return [str(x).lower() for x in parsed]
    # last resort: return words from job_text (simple heuristics)
    tokens = re.findall(r"[A-Za-z0-9+#\.]+", job_text)
    freq = {}
    for t in tokens:
        t2 = t.lower()
        if len(t2) > 2:
            freq[t2] = freq.get(t2, 0) + 1
    sorted_tokens = sorted(freq.items(), key=lambda x: -x[1])
    return [t for t, _ in sorted_tokens[:max_keywords]]


def extract_resume_bullets(llm, resume_text: str, max_bullets: int = 40) -> List[str]:
    prompt = (
        "Extract the main professional bullet points (experience, achievements, responsibilities) "
        "from the following resume text. Return a JSON array of strings, each string being a single bullet (no numbering).\n\nResume:\n"
    ) + resume_text

    messages = [
        ("system", "You are an assistant that extracts resume bullets in compact form."),
        ("human", prompt),
    ]
    ai_msg = llm.invoke(messages)
    content = getattr(ai_msg, 'content', str(ai_msg))
    parsed = safe_json_extract(content)
    if isinstance(parsed, list):
        bullets = [str(x).strip() for x in parsed if str(x).strip()]
        return bullets[:max_bullets]
    if isinstance(parsed, dict):
        # try to find list values
        for v in parsed.values():
            if isinstance(v, list):
                return [str(x).strip() for x in v][:max_bullets]
    # fallback: heuristics: break by lines and sentences
    lines = [l.strip(' -•\t') for l in resume_text.splitlines() if len(l.strip()) > 10]
    bullets = []
    for l in lines:
        if len(bullets) >= max_bullets:
            break
        if len(l) > 40 or ('%' in l) or re.search(r'\d{4}', l):
            bullets.append(l)
    # if nothing found, split by sentences
    if not bullets:
        sents = re.split(r'(?<=[\.!?])\s+', resume_text)
        for s in sents:
            s = s.strip()
            if len(s) > 40:
                bullets.append(s)
                if len(bullets) >= max_bullets:
                    break
    return bullets


def reword_variations_for_bullet(llm, bullet: str, n: int = 3) -> List[str]:
    prompt = (
        f"Rewrite this resume bullet in {n} different, concise and achievement-oriented ways. "
        "Return a JSON array of strings. Bullet:\n" + bullet
    )
    messages = [
        ("system", "You are an expert resume editor. Be concise and use action verbs and metrics where possible."),
        ("human", prompt),
    ]
    ai_msg = llm.invoke(messages)
    content = getattr(ai_msg, 'content', str(ai_msg))
    parsed = safe_json_extract(content)
    if isinstance(parsed, list):
        return [str(x).strip() for x in parsed if str(x).strip()][:n]
    # fallback: naive split by lines
    lines = [l.strip('-• ') for l in content.splitlines() if l.strip()]
    return lines[:n]

# -------------------- Embedding helpers --------------------

def cosine_sim(a: List[float], b: List[float]) -> float:
    a = np.array(a, dtype=float)
    b = np.array(b, dtype=float)
    if a.shape[0] == 0 or b.shape[0] == 0:
        return 0.0
    denom = (np.linalg.norm(a) * np.linalg.norm(b))
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)


def semantic_match_keyword_against_sentences(embeddings, keyword: str, sentences: List[str], threshold: float = 0.78) -> Tuple[bool, float, int]:
    """Return (matched_bool, best_score, best_sentence_index)"""
    try:
        kvec = embeddings.embed_query(keyword)
        svecs = embeddings.embed_documents(sentences)
    except Exception:
        return False, 0.0, -1
    best_score = 0.0
    best_idx = -1
    for i, v in enumerate(svecs):
        s = cosine_sim(kvec, v)
        if s > best_score:
            best_score = s
            best_idx = i
    return (best_score >= threshold), best_score, best_idx

# -------------------- UI / app flow --------------------



st.set_page_config(page_title='Agentic Resume Tailor', layout='wide')
st.title('Agentic Resume Tailor — Streamlit (LangChain + Gemini)')
st.caption('Prototype: LangChain for parsing + Gemini for rewording & embeddings')

with st.sidebar:
    st.header('Configuration')
    api_key_input = st.text_input('Google AI API Key (or set GOOGLE_API_KEY env var)', type='password')
    show_raw = st.checkbox('Show raw extracted text', value=False)
    sample_mode = st.checkbox('Use fast demo mode (fewer LLM calls)', value=False)

# let user upload files
col1, col2 = st.columns(2)
with col1:
    resume_file = st.file_uploader('Upload resume (PDF, DOCX, TXT)', type=['pdf', 'docx', 'txt'])
with col2:
    jd_file = st.file_uploader('Upload job description (PDF, DOCX, TXT)', type=['pdf', 'docx', 'txt'])

if st.button('Analyze'):
    if not resume_file or not jd_file:
        st.error('Please upload both a resume and a job description file.')
    else:
        # ensure API key available
        api_key = api_key_input.strip() or os.environ.get('GOOGLE_API_KEY')
        try:
            with st.spinner('Initializing LLM & embeddings...'):
                llm, embeddings = init_llm_and_embeddings(api_key)
        except Exception as e:
            st.error('LLM/Embeddings init error: %s' % str(e))
            st.stop()

        # extract text
        resume_text, resume_name = extract_text_from_uploaded(resume_file)
        jd_text, jd_name = extract_text_from_uploaded(jd_file)

        if show_raw:
            st.subheader('Raw extracted resume text')
            st.text_area('Resume raw', resume_text, height=300)
            st.subheader('Raw extracted job description text')
            st.text_area('JD raw', jd_text, height=300)

        # 1) extract JD keywords
        with st.spinner('Extracting keywords from job description...'):
            jd_keywords = extract_keywords_from_jd(llm, jd_text, max_keywords=50)
        st.success(f'Found {len(jd_keywords)} keywords (approx)')

        # 2) extract resume bullets
        with st.spinner('Extracting bullets from resume...'):
            resume_bullets = extract_resume_bullets(llm, resume_text, max_bullets=60)
        st.success(f'Extracted {len(resume_bullets)} resume bullets')

        # 3) prepare sentence chunks for semantic matching
        sentences = [s.strip() for s in re.split(r'\n|(?<=[\.!?])\s+', resume_text) if len(s.strip()) > 20]
        if not sentences:
            sentences = resume_bullets[:50]

        # 4) Match keywords (exact + semantic)
        matched = {}
        sem_threshold = 0.78
        with st.spinner('Matching keywords against resume...'):
            for kw in jd_keywords:
                kw_clean = kw.lower().strip()
                # exact match
                if re.search(r'\b' + re.escape(kw_clean) + r'\b', resume_text, flags=re.I):
                    matched[kw] = {'matched': True, 'method': 'exact', 'score': 1.0, 'sentence_idx': -1}
                    continue
                # semantic match via embeddings
                try:
                    ok, score, idx = semantic_match_keyword_against_sentences(embeddings, kw_clean, sentences, threshold=sem_threshold)
                except Exception:
                    ok, score, idx = False, 0.0, -1
                matched[kw] = {'matched': bool(ok), 'method': 'semantic' if ok else 'none', 'score': float(score), 'sentence_idx': int(idx)}

        # 5) Alignment score
        total = len(jd_keywords) if len(jd_keywords) > 0 else 1
        matched_count = sum(1 for v in matched.values() if v.get('matched'))
        alignment = int(matched_count / total * 100)
        if alignment < 50:
            color = 'red'
        elif alignment < 75:
            color = 'orange'
        else:
            color = 'green'

        st.markdown(f"**Alignment score:** <span style='color:{color}; font-weight:700'>{alignment}%</span>", unsafe_allow_html=True)

        # 6) Display keyword table with color-coded badges
        st.subheader('Job keywords — matched / missing')
        badge_lines = []
        for k, v in matched.items():
            if v['matched']:
                badge = f"<span style='background:#d4edda;color:#155724;border-radius:6px;padding:4px;margin:3px;display:inline-block'>{k}</span>"
            else:
                badge = f"<span style='background:#f8d7da;color:#721c24;border-radius:6px;padding:4px;margin:3px;display:inline-block'>{k}</span>"
            badge_lines.append(badge)
        st.markdown(' '.join(badge_lines), unsafe_allow_html=True)

        # 7) Highlight resume text with matched keywords
        st.subheader('Resume preview (highlighted)')
        highlighted = resume_text
        # avoid nested replacements by sorting by length desc
        kws_sorted = sorted(list(jd_keywords), key=lambda x: -len(x))
        for k in kws_sorted:
            if matched.get(k, {}).get('matched'):
                # wrap in green background
                highlighted = re.sub(r'(?i)\\b(' + re.escape(k) + r')\\b', r"<mark>\1</mark>", highlighted)
        st.markdown('<div style="white-space:pre-wrap;">' + highlighted + '</div>', unsafe_allow_html=True)

        # 8) For bullets that are related to missing keywords, offer reworded suggestions
        st.subheader('Tailored bullet suggestions')
        bullets_to_show = resume_bullets
        if sample_mode:
            bullets_to_show = resume_bullets[:6]

        new_bullets = []
        for i, b in enumerate(bullets_to_show):
            st.markdown(f'**Original bullet {i+1}**: {b}')
            # request variations
            with st.spinner(f'Generating variations for bullet {i+1}...'):
                try:
                    variations = reword_variations_for_bullet(llm, b, n=3)
                except Exception:
                    variations = []
            if variations:
                choice = st.selectbox(f'Pick a variation for bullet {i+1}', options=['(keep original)'] + variations, index=0, key=f'pick_{i}')
                if choice and choice != '(keep original)':
                    edited = st.text_area(f'Edit chosen bullet {i+1}', value=choice, key=f'edit_{i}')
                    new_bullets.append(edited)
                else:
                    edited = st.text_area(f'Edit original bullet {i+1}', value=b, key=f'editorig_{i}')
                    new_bullets.append(edited)
            else:
                edited = st.text_area(f'Edit original bullet {i+1}', value=b, key=f'edit_no_var_{i}')
                new_bullets.append(edited)

        # 9) Finalize + export
        st.subheader('Finalize & export')
        full_name = st.text_input('Name (for resume header)', value='')
        contact = st.text_input('Contact (email / phone)', value='')

        if st.button('Generate final DOCX and PDF'):
            # assemble a simple resume document from provided header and new_bullets
            docx_io = io.BytesIO()
            try:
                doc = docx.Document()
                if full_name:
                    doc.add_heading(full_name, level=1)
                if contact:
                    doc.add_paragraph(contact)
                doc.add_paragraph('')
                doc.add_heading('Experience', level=2)
                for b in new_bullets:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(b)
                doc.save(docx_io)
                docx_io.seek(0)
                st.success('DOCX created')

                # prepare PDF using reportlab
                pdf_io = io.BytesIO()
                elements = []
                styles = getSampleStyleSheet()
                if full_name:
                    elements.append(Paragraph(full_name, styles['Title']))
                if contact:
                    elements.append(Paragraph(contact, styles['Normal']))
                elements.append(Spacer(1, 12))
                elements.append(Paragraph('Experience', styles['Heading2']))
                for b in new_bullets:
                    elements.append(Paragraph(b, styles['Normal']))
                    elements.append(Spacer(1, 6))
                SimpleDocTemplate(pdf_io, pagesize=letter).build(elements)
                pdf_bytes = pdf_io.getvalue()

                st.download_button('Download DOCX', data=docx_io.getvalue(), file_name='tailored_resume.docx')
                st.download_button('Download PDF', data=pdf_bytes, file_name='tailored_resume.pdf')
            except Exception as e:
                st.error('Error creating files: %s' % str(e))

        # end analyze

# If user didn't press analyze show brief instructions
if not (resume_file and jd_file):
    st.info('Upload a resume and a job description then press Analyze.\n\nThis demo uses LangChain + langchain-google-genai (Gemini).')

# -------------------- END --------------------
